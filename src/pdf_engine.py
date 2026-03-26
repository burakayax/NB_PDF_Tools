import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageOps
import PyPDF2
import io
import os
import re
import statistics
import sys
import tempfile
from typing import List, Dict, Any, Tuple, Optional

# --- YOLLAR ---
# Tesseract'in yolu: sisteminizde farklıysa burayı güncelleyin
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
base_dir = os.path.dirname(os.path.abspath(__file__))
# Poppler klasörünün proje ana dizinindeki Library/bin içinde olduğu varsayılır
poppler_bin_path = os.path.join(base_dir, "..", "Library", "bin")

# Taranmış form / ekran görüntüsü PDF'ler için sayfa bölümleme (PSM 3: otomatik)
_TESSERACT_CONFIG = r"--oem 3 --psm 3"


def _open_pdf_reader(pdf_path: str, password: Optional[str] = None, context: str = "Bu işlem") -> PyPDF2.PdfReader:
    """PDF okuyucu oluşturur; parola verilmişse şifre çözümü yapar.
    Tüm sayfa işlemleri ortak açılış ve hata mesajı üretmek için buradan geçer.
    Parola politikası gevşetilirse şifreli dosyalar sessizce başarısız olabilir."""
    try:
        reader = PyPDF2.PdfReader(pdf_path)
        if reader.is_encrypted:
            if not password:
                raise Exception(
                    f"{context} için seçtiğiniz dosya şifreli: {os.path.basename(pdf_path)}\n"
                    "Lütfen dosya için şifre girin."
                )
            decrypt_result = reader.decrypt(password)
            if not decrypt_result:
                raise Exception(
                    f"{context} için girilen şifre hatalı: {os.path.basename(pdf_path)}"
                )
        return reader
    except Exception as e:
        err_text = str(e)
        if "şifreli" in err_text.lower() or "şifre" in err_text.lower():
            raise
        raise Exception(f"PDF dosyası kontrol edilemedi: {e}") from e


def _apply_output_pdf_password(output_path: str, output_password: Optional[str]) -> None:
    """Yazılmış PDF çıktısına görüntüleyici parolası uygular (pikepdf ile).
    Parola boşsa dosyaya dokunmaz; şifreleme modülü akışını sadeleştirir.
    pikepdf yoksa veya dosya kilitliyse çağıran kullanıcıya anlamlı hata göstermelidir."""
    if not output_password:
        return
    try:
        import pikepdf
    except ImportError as e:
        raise Exception("PDF çıktı şifreleme için 'pikepdf' gerekli.") from e

    owner = output_password
    out_dir = os.path.dirname(output_path) or None
    fd, temp_output_path = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        with pikepdf.open(output_path) as pdf:
            pdf.save(
                temp_output_path,
                encryption=pikepdf.Encryption(
                    user=output_password,
                    owner=owner,
                    R=6,
                    allow=pikepdf.Permissions(extract=False),
                ),
            )
        os.replace(temp_output_path, output_path)
    finally:
        if os.path.exists(temp_output_path):
            try:
                os.remove(temp_output_path)
            except Exception:
                pass


def is_pdf_encrypted(pdf_path: str) -> bool:
    """Verilen PDF şifreliyse True döndürür.
    Arayüzde parola alanını göstermek için hızlı kontrol sağlar.
    Dosya bozuksa veya okunamazsa istisna yükseltilir."""
    try:
        with open(pdf_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            return bool(reader.is_encrypted)
    except Exception as e:
        raise Exception(f"PDF dosyası kontrol edilemedi: {e}") from e


def validate_pdf_password(pdf_path: str, password: str) -> bool:
    """Parola PDF'i açabiliyorsa True döndürür; şifresiz dosyada her zaman True.
    Şifreli dosyada yanlış parola False üretir.
    Okuma hatalarında açıklamalı istisna fırlatılır."""
    try:
        reader = PyPDF2.PdfReader(pdf_path)
        if not reader.is_encrypted:
            return True
        return bool(reader.decrypt(password or ""))
    except Exception as e:
        raise Exception(f"PDF şifresi doğrulanamadı: {e}") from e


def get_num_pages(pdf_path: str, password: Optional[str] = None) -> int:
    """PDF içindeki sayfa sayısını döndürür; şifreliyse parola ile açar.
    Bölme ve önizleme akışları sayfa sınırı bilmek zorundadır.
    Parola eksik veya hatalıysa _open_pdf_reader üzerinden anlamlı hata verilir."""
    try:
        reader = _open_pdf_reader(pdf_path, password=password, context="Sayfa bilgisi okuma")
        return len(reader.pages)
    except Exception as e:
        raise Exception(f"PDF sayfa sayısı okunamadı: {e}")


def _word_to_pdf_win32com(docx_path: str, pdf_path: str) -> None:
    """
    Microsoft Word COM (pywin32) ile PDF dışa aktarır.
    Arka plan iş parçacığında çağrılıyorsa CoInitialize gerekir.
    """
    try:
        import pythoncom
        from win32com.client import DispatchEx
    except ImportError as e:
        raise ImportError("pywin32 paketi yüklü değil.") from e

    pythoncom.CoInitialize()
    word = None
    try:
        word = DispatchEx("Word.Application")
        word.Visible = False
        try:
            word.DisplayAlerts = 0
        except Exception:
            pass
        doc_abs = os.path.abspath(docx_path)
        out_abs = os.path.abspath(pdf_path)
        doc = word.Documents.Open(doc_abs, ReadOnly=True)
        # Word sabiti: wdExportFormatPDF = 17 (sabit sayı COM API ile uyumludur).
        doc.ExportAsFixedFormat(OutputFileName=out_abs, ExportFormat=17, OpenAfterExport=False)
        doc.Close(SaveChanges=False)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def word_to_pdf(docx_path: str, pdf_path: str, progress_callback=None) -> bool:
    """
    Word belgesini PDF'e çevirir.

    Windows: Microsoft Word gerekir. Önce docx2pdf, yoksa pywin32 (COM) denenir.
    macOS: Microsoft Word + docx2pdf.
    """
    if sys.platform not in ("win32", "darwin"):
        raise Exception(
            "Word'den PDF dönüşümü şu an yalnızca Windows veya macOS üzerinde "
            "Microsoft Word yüklüyken desteklenir."
        )
    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"Dosya bulunamadı: {docx_path}")
    ext = os.path.splitext(docx_path)[1].lower()
    if ext not in (".docx", ".doc"):
        raise ValueError("Yalnızca .docx veya .doc dosyaları desteklenir.")

    pdf_path = os.path.abspath(pdf_path)
    out_dir = os.path.dirname(pdf_path)
    if out_dir and not os.path.isdir(out_dir):
        raise FileNotFoundError(f"Hedef klasör bulunamadı: {out_dir}")

    if progress_callback:
        progress_callback(0, 2, "Microsoft Word ile PDF oluşturuluyor...")

    last_err: Optional[Exception] = None
    if sys.platform == "win32":
        try:
            from docx2pdf import convert

            try:
                convert(docx_path, pdf_path)
            except Exception as e:
                last_err = e
                try:
                    _word_to_pdf_win32com(docx_path, pdf_path)
                except Exception as e2:
                    raise Exception(
                        f"Word -> PDF Hatası (docx2pdf): {last_err}\n"
                        f"Yedek COM yolu da başarısız: {e2}"
                    ) from e2
        except ImportError:
            try:
                _word_to_pdf_win32com(docx_path, pdf_path)
            except ImportError:
                raise Exception(
                    "Word'den PDF için paket gerekli. PowerShell veya CMD'de (uygulamayı çalıştırdığın Python ile):\n\n"
                    "  python -m pip install docx2pdf\n\n"
                    "veya sadece COM yolu için:\n\n"
                    "  python -m pip install pywin32\n\n"
                    "Microsoft Word kurulu olmalıdır."
                ) from None
            except Exception as e:
                raise Exception(f"Word -> PDF Hatası: {e}") from e
    else:
        try:
            from docx2pdf import convert

            convert(docx_path, pdf_path)
        except ImportError as e:
            raise Exception(
                "Word'den PDF için 'docx2pdf' paketi gerekli. Kurulum:\n"
                "python -m pip install docx2pdf\n\n"
                "Microsoft Word for Mac kurulu olmalıdır."
            ) from e
        except Exception as e:
            raise Exception(f"Word -> PDF Hatası: {e}") from e

    if not os.path.isfile(pdf_path):
        raise Exception(
            "PDF dosyası oluşmadı. Word kapalı olsun, dosya başka programda açık olmasın "
            "veya Word güvenlik uyarısı engellenmiş olabilir."
        )

    if progress_callback:
        progress_callback(2, 2, "Tamamlandı")
    return True


def _fitz_pdf_text_stats(pdf_path: str, password: Optional[str]) -> Tuple[int, int]:
    """Sayfa sayısı ve çıkarılabilir metin uzunluğu (taranmış PDF ayırt etmek için)."""
    import fitz

    doc = fitz.open(pdf_path)
    try:
        if doc.needs_pass:
            if not password:
                raise Exception("Şifre gerekli")
            if not doc.authenticate(password):
                raise Exception("Girilen şifre hatalı")
        n = len(doc)
        total = 0
        for i in range(n):
            total += len(doc.load_page(i).get_text())
        return n, total
    finally:
        doc.close()


def _pdf_to_word_ocr_fitz(
    pdf_path: str,
    docx_path: str,
    password: Optional[str],
    progress_callback=None,
) -> None:
    """Metin katmanı zayıf / taranmış PDF: sayfa görüntüsünden Tesseract ile düzenlenebilir metin."""
    import fitz

    doc_pdf = fitz.open(pdf_path)
    try:
        if doc_pdf.needs_pass:
            if not password:
                raise Exception("Şifre gerekli")
            if not doc_pdf.authenticate(password):
                raise Exception("Girilen şifre hatalı")
        dw = Document()
        n = len(doc_pdf)
        for i in range(n):
            if progress_callback:
                progress_callback(i + 1, max(n, 1), f"OCR (sayfa {i + 1}/{n})")
            page = doc_pdf.load_page(i)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            try:
                text = pytesseract.image_to_string(img, lang="tur+eng", config=_TESSERACT_CONFIG)
            except Exception:
                text = pytesseract.image_to_string(img, lang="eng", config=_TESSERACT_CONFIG)
            text = _polish_tesseract_output(text.strip())
            if i > 0:
                dw.add_page_break()
            dw.add_heading(f"Sayfa {i + 1}", level=2)
            if text:
                for block in text.split("\n\n"):
                    line = block.strip()
                    if line:
                        dw.add_paragraph(line)
            else:
                dw.add_paragraph("(Bu sayfada metin okunamadı.)")
        dw.save(docx_path)
    finally:
        doc_pdf.close()


def pdf_to_word(pdf_path: str, docx_path: str, progress_callback=None, password: Optional[str] = None) -> bool:
    """
    PDF'i düzenlenebilir DOCX olarak dönüştürür.
    Metin katmanı çok zayıfsa Tesseract OCR ile metin üretir; aksi halde pdf2docx (yapısal) kullanılır.
    """
    try:
        if is_pdf_encrypted(pdf_path) and not password:
            raise Exception(f"PDF'ten Word'e dönüşüm için şifre gerekli: {os.path.basename(pdf_path)}")

        if progress_callback:
            progress_callback(0, 3, "PDF analiz ediliyor...")

        try:
            from pdf2docx import Converter
        except ImportError as e:
            raise Exception("pdf2docx yüklü değil.") from e

        n_pages, text_len = _fitz_pdf_text_stats(pdf_path, password)
        # Taranmış PDF: yalnızca gerçekten metin katmanı çok zayıfsa OCR (aksi halde pdf2docx düzen korur; erken OCR bazı metin PDF’lerini bozuyordu)
        if n_pages > 0 and text_len < max(50, n_pages * 18):
            if progress_callback:
                progress_callback(1, 3, "Taranmış veya görsel PDF — OCR ile metin çıkarılıyor...")
            if os.path.isfile(docx_path):
                try:
                    os.remove(docx_path)
                except OSError:
                    pass
            _pdf_to_word_ocr_fitz(pdf_path, docx_path, password, progress_callback=progress_callback)
            if progress_callback:
                progress_callback(3, 3, "Word kaydediliyor...")
            return True

        if progress_callback:
            progress_callback(0, 3, "Yapısal dönüşüm hazırlanıyor...")

        def _run_convert(use_ocr: int) -> None:
            converter = Converter(pdf_path, password=password)
            try:
                if progress_callback:
                    msg = (
                        "Taranmış sayfalar için OCR ile Word oluşturuluyor..."
                        if use_ocr
                        else "Sayfa düzeni korunarak Word oluşturuluyor..."
                    )
                    progress_callback(1, 3, msg)
                if os.path.isfile(docx_path):
                    try:
                        os.remove(docx_path)
                    except OSError:
                        pass
                converter.convert(
                    docx_path,
                    ocr=use_ocr,
                    clip_image_res_ratio=2.0,
                    float_image_ignorable_gap=12.0,
                    line_overlap_threshold=0.9,
                    line_separate_threshold=6.0,
                )
            finally:
                converter.close()

        try:
            _run_convert(0)
        except Exception as structural_error:
            try:
                _run_convert(1)
            except Exception:
                try:
                    if os.path.isfile(docx_path):
                        try:
                            os.remove(docx_path)
                        except OSError:
                            pass
                    if progress_callback:
                        progress_callback(2, 3, "Yapısal dönüşüm başarısız — OCR deneniyor...")
                    _pdf_to_word_ocr_fitz(pdf_path, docx_path, password, progress_callback=progress_callback)
                except Exception:
                    raise Exception(
                        "Bu PDF düzenlenebilir Word olarak dönüştürülemedi. "
                        "Tesseract kurulu olmalı; taranmış PDF’lerde OCR gerekir.\n\n"
                        f"Teknik ayrıntı: {structural_error}"
                    ) from structural_error

        if os.path.isfile(docx_path):
            if progress_callback:
                progress_callback(3, 3, "Word kaydediliyor...")
            return True

        raise Exception("Word dosyası oluşturulamadı.")

    except Exception as e:
        raise Exception(f"PDF -> Word Hatası: {e}")


def _safe_int(v, default=0) -> int:
    try:
        return int(v)
    except Exception:
        return default


def _safe_float(v, default=-1.0) -> float:
    try:
        return float(v)
    except Exception:
        return default


def _run_tesseract_image_to_data(page_proc: Image.Image) -> Dict[str, Any]:
    """
    Önce sadece Türkçe (Ü, Ğ vb. için daha iyi); yoksa tur+eng ile dener.
    """
    kwargs = dict(output_type=pytesseract.Output.DICT, config=_TESSERACT_CONFIG)
    try:
        return pytesseract.image_to_data(page_proc, lang="tur", **kwargs)
    except Exception:
        return pytesseract.image_to_data(page_proc, lang="tur+eng", **kwargs)


def _polish_tesseract_output(s: str) -> str:
    """
    Tesseract'ın sık yaptığı Türkçe / rakam hatalarını metin düzeyinde düzeltir.
    """
    if not s:
        return s
    # Ürün (Ü harfi bazen Uriin / Urin olarak gelir)
    s = re.sub(r"Uriin", "Ürün", s, flags=re.IGNORECASE)
    s = re.sub(r"\bUrin\b", "Ürün", s, flags=re.IGNORECASE)
    s = re.sub(r"\bUrun\b", "Ürün", s, flags=re.IGNORECASE)
    # Bitişik başlık
    s = re.sub(r"Ürün\s*Takip\s*Sistemi", "Ürün Takip Sistemi", s, flags=re.IGNORECASE)
    s = re.sub(r"ÜrünTakipSistemi", "Ürün Takip Sistemi", s, flags=re.IGNORECASE)
    # Rakamlar arası þ (thorn) -> 0 (ör. 735þ658411)
    s = re.sub(r"(\d)[þÞ](\d)", r"\g<1>0\g<2>", s)
    # Sadece rakam benzeri parçalarda kalan þ
    def fix_token(tok: str) -> str:
        if re.search(r"\d", tok) and re.search(r"[þÞ]", tok):
            return tok.replace("þ", "0").replace("Þ", "0")
        return tok

    parts = re.split(r"(\s+)", s)
    s = "".join(fix_token(p) if p.strip() else p for p in parts)
    return re.sub(r" {2,}", " ", s).strip()


def _finalize_line_text(s: str) -> str:
    return _polish_tesseract_output(_fix_glued_turkish_text(s))


def _fix_glued_turkish_text(s: str) -> str:
    """OCR/Tesseract bazen kelimeleri bitişik yazar (ör. FirmaAdi). Metin düzenlenebilir kalsın diye ayırır."""
    if not s:
        return s
    # küçük harf + büyük harf ayrımı (camelCase benzeri)
    s = re.sub(r"([a-zığüşöç])([A-ZĞÜŞİÖÇİ])", r"\1 \2", s)
    # rakam <-> harf
    s = re.sub(r"(\d)([A-Za-zğüşıöçĞÜŞİÖÇİ])", r"\1 \2", s)
    s = re.sub(r"([A-Za-zğüşıöçĞÜŞİÖÇİ])(\d)", r"\1 \2", s)
    # yaygın form etiketleri (ÜTS / bayilik ekranları)
    replacements = (
        ("FirmaAdi", "Firma Adı"),
        ("FirmaVergi", "Firma Vergi"),
        ("BayilikVeren", "Bayilik Veren"),
        ("BayilikAlan", "Bayilik Alan"),
        ("BayilikBaşvuru", "Bayilik Başvuru"),
        ("BayilikBasvuru", "Bayilik Başvuru"),
        ("BaşvuruTarihi", "Başvuru Tarihi"),
        ("BasvuruTarihi", "Başvuru Tarihi"),
        ("BaşlangıçTarihi", "Başlangıç Tarihi"),
        ("BaslangicTarihi", "Başlangıç Tarihi"),
        ("PlanlananBitiş", "Planlanan Bitiş"),
        ("PlanlananBitis", "Planlanan Bitiş"),
        ("BitişTarihi", "Bitiş Tarihi"),
        ("BitisTarihi", "Bitiş Tarihi"),
        ("KararTarihi", "Karar Tarihi"),
        ("İthalatBildirimi", "İthalat Bildirimi"),
        ("IthalatBildirimi", "İthalat Bildirimi"),
    )
    for a, b in replacements:
        s = s.replace(a, b)
    return re.sub(r" {2,}", " ", s).strip()


def _parse_ocr_words(ocr: Dict[str, Any]) -> List[Dict[str, Any]]:
    texts = ocr.get("text", [])
    if not texts:
        return []

    lefts = ocr.get("left", [])
    tops = ocr.get("top", [])
    widths = ocr.get("width", [])
    heights = ocr.get("height", [])
    block_nums = ocr.get("block_num", [0] * len(texts))
    par_nums = ocr.get("par_num", [0] * len(texts))
    line_nums = ocr.get("line_num", [0] * len(texts))
    confs = ocr.get("conf", [-1] * len(texts))

    words: List[Dict[str, Any]] = []
    n = len(texts)
    for i in range(n):
        t = (texts[i] or "").strip()
        if not t:
            continue
        conf = _safe_float(confs[i], default=-1.0)
        if conf < 0:
            conf = 0.0
        words.append(
            {
                "text": t,
                "left": _safe_int(lefts[i] if i < len(lefts) else 0),
                "top": _safe_int(tops[i] if i < len(tops) else 0),
                "width": _safe_int(widths[i] if i < len(widths) else 0),
                "height": _safe_int(heights[i] if i < len(heights) else 0),
                "block": _safe_int(block_nums[i] if i < len(block_nums) else 0),
                "par": _safe_int(par_nums[i] if i < len(par_nums) else 0),
                "line": _safe_int(line_nums[i] if i < len(line_nums) else 0),
                "conf": conf,
            }
        )
    return words


def _cluster_words_into_visual_lines(words: List[Dict[str, Any]]) -> List[List[Dict[str, Any]]]:
    """Tüm blokları birleştirip y koordinatına göre görsel satırlara ayırır."""
    if not words:
        return []
    heights = sorted(w["height"] for w in words if w["height"] > 0)
    median_h = heights[len(heights) // 2] if heights else 14
    y_tol = max(10, int(median_h * 0.52))

    sorted_w = sorted(words, key=lambda w: (w["top"], w["left"]))
    lines: List[List[Dict[str, Any]]] = []
    current: List[Dict[str, Any]] = []
    anchor_y: float = 0.0

    for w in sorted_w:
        cy = float(w["top"] + max(1, w["height"]) / 2.0)
        if not current:
            current = [w]
            anchor_y = cy
            continue
        if abs(cy - anchor_y) <= y_tol:
            current.append(w)
            anchor_y = (anchor_y * (len(current) - 1) + cy) / len(current)
        else:
            current.sort(key=lambda x: x["left"])
            lines.append(current)
            current = [w]
            anchor_y = cy

    if current:
        current.sort(key=lambda x: x["left"])
        lines.append(current)
    return lines


def _join_words_in_segment(words: List[Dict[str, Any]]) -> str:
    if not words:
        return ""
    words = sorted(words, key=lambda x: x["left"])
    hlist = [w["height"] for w in words if w["height"] > 0]
    med_h = statistics.median(hlist) if hlist else 12.0
    space_threshold = max(2.0, med_h * 0.11)

    parts: List[str] = []
    for i, w in enumerate(words):
        if i == 0:
            parts.append(w["text"])
            continue
        prev = words[i - 1]
        gap = float(w["left"] - (prev["left"] + prev["width"]))
        sep = " " if gap > space_threshold else ""
        parts.append(sep + w["text"])
    raw = "".join(parts)
    return _finalize_line_text(raw.replace("  ", " ").strip())


def _split_line_into_word_segments(words: List[Dict[str, Any]], page_w: int) -> List[List[Dict[str, Any]]]:
    """Satır içinde etiket / değer gibi geniş boşluklardan kelime grupları üretir."""
    words = sorted(words, key=lambda x: x["left"])
    if len(words) <= 1:
        return [words]

    gaps: List[float] = []
    for i in range(1, len(words)):
        g = float(words[i]["left"] - (words[i - 1]["left"] + words[i - 1]["width"]))
        gaps.append(max(0.0, g))

    med_g = statistics.median(gaps) if gaps else 0.0
    # Form satırlarında etiket-değer arası genelde çok geniş; küçük iç boşluklara dokunma
    threshold = max(42.0, page_w * 0.052, med_g * 2.6 if med_g > 4.0 else 42.0)

    segments: List[List[Dict[str, Any]]] = []
    start = 0
    for i, g in enumerate(gaps):
        if g >= threshold:
            segments.append(words[start : i + 1])
            start = i + 1
    segments.append(words[start:])
    return [s for s in segments if s]


def _line_bbox(words: List[Dict[str, Any]]) -> Tuple[int, int]:
    left = min(w["left"] for w in words)
    right = max(w["left"] + w["width"] for w in words)
    return left, right


def _line_looks_centered(words: List[Dict[str, Any]], page_w: int) -> bool:
    if len(words) < 1 or page_w <= 0:
        return False
    left, right = _line_bbox(words)
    span = right - left
    mid = (left + right) / 2.0
    return abs(mid - page_w / 2.0) < page_w * 0.10 and span < page_w * 0.78


def _try_split_label_value_text(t: str) -> Optional[Tuple[str, str]]:
    """
    'Telefon: 0(212)...' veya 'Firma Adı: ...' gibi tek parçada kalan satırları [etiket, değer] ayırır.
    """
    t = (t or "").strip()
    if not t or t.lower().startswith("http"):
        return None
    first = t.split("\n", 1)[0].strip()
    if ":" not in first:
        return None
    idx = first.index(":")
    label = first[:idx].strip()
    rest = t[idx + 1 :].strip()
    if len(label) < 2 or len(label) > 58:
        return None
    # Çok geniş eşleşmeleri engelle (cümle içi iki nokta üst üste)
    if "\n" in label:
        return None
    return (label + ":", rest)


def _try_paragraph_to_two_column_row(text: str) -> Optional[Dict[str, Any]]:
    pair = _try_split_label_value_text(text)
    if not pair:
        return None
    label, value = pair
    return {"kind": "table_row", "cells": [label, value]}


def _looks_like_new_field_paragraph(text: str) -> bool:
    """Yeni bir form alanı satırı mı (sol tarafta etiket + iki nokta)?"""
    first = text.strip().split("\n", 1)[0].strip()
    if ":" not in first:
        return False
    idx = first.index(":")
    if idx > 58 or idx < 2:
        return False
    before = first[:idx].strip()
    if len(before) > 55:
        return False
    return True


def _looks_like_firma_adi_continuation(text: str) -> bool:
    """Firma adı değerinin ikinci satırı (SAN. TİC. LTD. ŞTİ. vb.)"""
    t = text.strip()
    if not t:
        return False
    if _looks_like_new_field_paragraph(t):
        return False
    if re.search(
        r"SAN\.\s*TİC\.|TİC\.\s*LTD|LTD\.?\s*ŞTİ\.?|ŞTİ\.?\s*\(|Üretici\s*/\s*İthalatçı",
        t,
        re.IGNORECASE,
    ):
        return True
    if t.startswith("(") and "Üretici" in t:
        return True
    return False


def _looks_like_phone_line(text: str) -> bool:
    """Telefon numarasına benzeyen kısa satır (OCR bazen etiketten ayırır)."""
    t = text.strip()
    if len(t) < 8:
        return False
    compact = re.sub(r"\s+", "", t)
    digits = sum(c.isdigit() for c in compact)
    if digits >= 10 and digits / max(1, len(compact)) > 0.5:
        return True
    if re.match(r"^0[\d\s\(\)\-]{8,}$", t.strip()):
        return True
    return False


def _should_merge_value_continuation(prev_row: Dict[str, Any], para_text: str) -> bool:
    """Önceki 2 sütunlu satırın sağ hücresine bu paragrafı eklemeli miyiz?"""
    if prev_row.get("kind") != "table_row":
        return False
    cells = prev_row.get("cells") or []
    if len(cells) != 2:
        return False
    left_l = cells[0].strip().lower()
    right = cells[1].strip()
    t = para_text.strip()
    if not t:
        return False
    if "bilgileri" in t.lower() and len(t) < 100 and "firma" not in t.lower():
        return False
    # Firma adı: uzun isim iki satıra bölünmüş
    if "firma adı" in left_l or "firma adi" in left_l:
        if _looks_like_firma_adi_continuation(t):
            return True
        if ":" not in t.split("\n")[0] and len(t) < 200:
            if right and not right.endswith(")"):
                if re.search(r"ŞTİ|TİC\.|LTD|SAN\.", t, re.IGNORECASE):
                    return True
    return False


def _normalize_single_cell_table_row(block: Dict[str, Any]) -> Dict[str, Any]:
    """Tek hücrede kalmış 'Etiket: değer' satırlarını ikiye böler."""
    if block.get("kind") != "table_row":
        return block
    cells = block.get("cells") or []
    if len(cells) != 1:
        return block
    pair = _try_split_label_value_text(cells[0])
    if not pair:
        return block
    return {"kind": "table_row", "cells": [pair[0], pair[1]]}


def _postprocess_form_layout_blocks(blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    - Paragraf olarak kalan 'Etiket: değer' satırlarını tablo satırına çevirir (Telefon hizası için).
    - Firma adı gibi alanlarda ikinci görsel satırı önceki satırın sağ hücresine birleştirir.
    """
    # A: paragrafları etiket/değer tablosuna çevir + tek hücreli tabloları böl
    pass1: List[Dict[str, Any]] = []
    for b in blocks:
        b = _normalize_single_cell_table_row(b)
        if b.get("kind") == "paragraph":
            conv = _try_paragraph_to_two_column_row(b.get("text", ""))
            if conv:
                pass1.append(conv)
                continue
        pass1.append(b)

    # B: devam satırlarını önceki satırın sağ hücresine yapıştır
    pass2: List[Dict[str, Any]] = []
    for b in pass1:
        if b.get("kind") == "paragraph" and pass2:
            prev = pass2[-1]
            if prev.get("kind") == "table_row" and len(prev.get("cells", [])) == 2:
                cells = prev["cells"]
                c0, c1 = cells[0].strip(), cells[1].strip()
                # Telefon etiketi tek satırda kaldıysa ve numara sonraki satırda geldiyse
                if (
                    c0.lower().startswith("telefon")
                    and ":" in c0
                    and _looks_like_phone_line(b.get("text", ""))
                    and len(c1) < 4
                ):
                    prev["cells"] = [cells[0], (c1 + " " + b["text"].strip()).strip()]
                    continue
            if _should_merge_value_continuation(prev, b.get("text", "")):
                cells = prev["cells"]
                prev["cells"] = [cells[0], (cells[1] + " " + b["text"].strip()).strip()]
                continue
        pass2.append(b)
    return pass2


def _is_section_heading_text(text: str) -> bool:
    t = text.strip()
    if len(t) > 90:
        return False
    if "Bilgileri" in t and len(t.split()) <= 8:
        return True
    if t.endswith(":") and len(t) < 50:
        return True
    return False


def _ocr_data_to_layout_blocks(ocr: Dict[str, Any], page_w: int) -> List[Dict[str, Any]]:
    """
    OCR çıktısını üstten alta sıralı blok listesine çevirir.
    - Etiket / değer: geniş boşlukla ayrılmış satırlar -> kenarlıksız tablo satırı
    - Tek sütun, ortada dar blok -> ortalanmış paragraf (başlık)
    """
    words = _parse_ocr_words(ocr)
    if not words:
        return []

    lines = _cluster_words_into_visual_lines(words)
    blocks: List[Dict[str, Any]] = []

    for line_words in lines:
        segs = _split_line_into_word_segments(line_words, page_w)
        texts = [_join_words_in_segment(s) for s in segs]
        texts = [t for t in texts if t.strip()]
        if not texts:
            continue

        if len(texts) == 1:
            text = texts[0]
            centered = _line_looks_centered(line_words, page_w)
            bold = _is_section_heading_text(text)
            blocks.append(
                {
                    "kind": "paragraph",
                    "text": text,
                    "centered": centered,
                    "bold": bold,
                }
            )
        else:
            # Çok sütun: üst bilgi (tarih | başlık | kullanıcı) veya yan yana alanlar
            blocks.append({"kind": "table_row", "cells": texts})

    return _postprocess_form_layout_blocks(blocks)


def _set_run_language_turkish(run) -> None:
    """Word yazım denetiminin Türkçe ile daha iyi çalışması için parça dilini ayarlar."""
    try:
        rpr = run._element.get_or_add_rPr()
        lang_el = OxmlElement("w:lang")
        lang_el.set(qn("w:val"), "tr-TR")
        lang_el.set(qn("w:eastAsia"), "tr-TR")
        rpr.append(lang_el)
    except Exception:
        pass


def _set_table_no_border(table) -> None:
    """Word tablosunda çizgileri kaldırır (düzen form görünümü)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def _append_layout_block_to_document(doc: Document, block: Dict[str, Any]) -> None:
    if block.get("kind") == "paragraph":
        text = (block.get("text") or "").strip()
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(text)
        run.bold = bool(block.get("bold"))
        _set_run_language_turkish(run)
        if block.get("centered"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    if block.get("kind") == "table_row":
        cells = [c.strip() for c in block.get("cells") or [] if str(c).strip()]
        if not cells:
            return
        n = len(cells)
        tbl = doc.add_table(rows=1, cols=n)
        _set_table_no_border(tbl)
        row = tbl.rows[0]
        for j, ctext in enumerate(cells):
            cell = row.cells[j]
            cell.text = ""
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.left_indent = Pt(0)
            r = cp.add_run(ctext)
            r.bold = False
            _set_run_language_turkish(r)
        if n == 3:
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # İki sütun: sol etiket, sağ değer (tipik form)
        if n == 2:
            try:
                row.cells[0].width = Inches(2.25)
                row.cells[1].width = Inches(4.35)
            except Exception:
                pass
        return


# --- 2. PDF BİRLEŞTİRME ---
def merge_pdfs(pdf_list: List[str], output_path: str, progress_callback=None, passwords: Optional[Dict[str, str]] = None) -> bool:
    """
    Birden fazla PDF dosyasını tek çıktıda birleştirir.

    İsteğe bağlı progress_callback çağrı biçimi:
        progress_callback(current: int, total: int, where_text: str) -> bool|None
    False dönerse birleştirme iptal edilir.
    """
    try:
        writer = PyPDF2.PdfWriter()
        readers = []
        total_pages = 0

        # Önce toplam sayfa sayısını toplayıp ilerleme çubuğunu dosya yerine sayfa bazlı akıtıyoruz.
        for pdf in pdf_list:
            if not os.path.isfile(pdf):
                raise FileNotFoundError(f"Birleştirilecek dosya bulunamadı: {pdf}")
            reader = _open_pdf_reader(pdf, password=(passwords or {}).get(pdf), context="PDF birleştirme")
            readers.append((pdf, reader))
            total_pages += len(reader.pages)

        current_page = 0
        total = max(1, total_pages + 1)
        for pdf, reader in readers:
            base_name = os.path.basename(pdf)
            page_count = len(reader.pages)
            for page_idx, page in enumerate(reader.pages, start=1):
                current_page += 1
                writer.add_page(page)
                if progress_callback:
                    res = progress_callback(
                        current_page,
                        total,
                        f"{base_name} | Sayfa {page_idx}/{page_count}",
                    )
                    if res is False:
                        raise Exception("İşlem iptal edildi.")

        if progress_callback:
            progress_callback(total, total, "PDF yazılıyor...")
        with open(output_path, "wb") as out_file:
            writer.write(out_file)
        return True
    except Exception as e:
        err_text = str(e)
        if "PyCryptodome is required for AES algorithm" in err_text:
            raise Exception(
                "Bazı PDF dosyaları AES şifreleme kullanıyor. Birleştirme için eksik paket bulundu:\n\n"
                "python -m pip install pycryptodome"
            ) from e
        raise Exception(f"Birleştirme Hatası: {e}")


# --- 3. SAYFA AYIKLA (tek PDF olarak) ---
def extract_pages(
    pdf_path: str,
    pages: List[int],
    output_path: str,
    password: Optional[str] = None,
    output_password: Optional[str] = None,
) -> bool:
    """
    pdf_path içinden verilen sayfaları (1 tabanlı liste) alır ve tek PDF olarak output_path'e yazar.
    Geçersiz sayfa numaralarında ValueError, diğer hatalarda Exception yükseltilir.
    Çıktı parolası isteniyorsa yazım sonrası _apply_output_pdf_password uygulanır.
    """
    try:
        reader = _open_pdf_reader(pdf_path, password=password, context="Sayfa ayıklama")
        num_pages = len(reader.pages)

        # Sayfa numaralarını doğrular (1 tabanlı indeks beklenir).
        normalized = []
        for p in pages:
            if not isinstance(p, int):
                raise ValueError(f"Sayfa numarası tam sayı olmalıdır: {p}")
            if p < 1 or p > num_pages:
                raise ValueError(f"Geçersiz sayfa numarası: {p} (Dosya {num_pages} sayfa)")
            normalized.append(p)

        writer = PyPDF2.PdfWriter()
        for p in normalized:
            writer.add_page(reader.pages[p - 1])

        with open(output_path, "wb") as f:
            writer.write(f)
        _apply_output_pdf_password(output_path, output_password)

        return True
    except Exception as e:
        raise Exception(f"Ayıklama Hatası: {e}")


# --- 4. SAYFA AYIKLA (her sayfa ayrı dosya olarak) ---
def extract_pages_separate(
    pdf_path: str,
    pages: List[int],
    output_folder: str,
    password: Optional[str] = None,
    output_password: Optional[str] = None,
) -> List[str]:
    """
    Verilen sayfaları (1 tabanlı) ayırır; her birini output_folder içinde ayrı PDF olarak kaydeder.
    Oluşan dosya yollarının listesini döndürür; hata durumunda istisna fırlatır.
    Klasör yoksa FileNotFoundError ile erken çıkış yapılır.
    """
    try:
        if not os.path.isdir(output_folder):
            raise FileNotFoundError(f"Hedef klasör bulunamadı: {output_folder}")

        reader = _open_pdf_reader(pdf_path, password=password, context="Sayfa ayıklama")
        num_pages = len(reader.pages)

        normalized = []
        for p in pages:
            if not isinstance(p, int):
                raise ValueError(f"Sayfa numarası tam sayı olmalıdır: {p}")
            if p < 1 or p > num_pages:
                raise ValueError(f"Geçersiz sayfa numarası: {p} (Dosya {num_pages} sayfa)")
            normalized.append(p)

        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        output_paths = []

        for p in normalized:
            writer = PyPDF2.PdfWriter()
            writer.add_page(reader.pages[p - 1])
            out_name = f"{base_name}_page_{p}.pdf"
            out_path = os.path.join(output_folder, out_name)
            with open(out_path, "wb") as f:
                writer.write(f)
            _apply_output_pdf_password(out_path, output_password)
            output_paths.append(out_path)

        return output_paths

    except Exception as e:
        raise Exception(f"Ayrı Ayırma Hatası: {e}")


# --- 5. PDF METİN / TABLO -> EXCEL ---
def _sanitize_sheet_title(title: str) -> str:
    invalid = '[]:*?/\\'
    cleaned = "".join("_" if ch in invalid else ch for ch in title).strip()
    return (cleaned or "Sayfa")[:31]


def _pdf_tables_to_excel(pdf_path: str, xlsx_path: str, progress_callback=None, password: Optional[str] = None) -> bool:
    try:
        import pdfplumber
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        raise Exception(
            "Tablo koruma modu için 'pdfplumber' ve 'openpyxl' gerekli: python -m pip install pdfplumber openpyxl"
        ) from e

    try:
        wb = Workbook()
        default_ws = wb.active
        wb.remove(default_ws)

        if is_pdf_encrypted(pdf_path) and not password:
            raise Exception(f"PDF -> Excel için şifre gerekli: {os.path.basename(pdf_path)}")
        with pdfplumber.open(pdf_path, password=password) as pdf:
            total_pages = len(pdf.pages)
            border = Border(
                left=Side(style="thin", color="D7DEE8"),
                right=Side(style="thin", color="D7DEE8"),
                top=Side(style="thin", color="D7DEE8"),
                bottom=Side(style="thin", color="D7DEE8"),
            )
            title_fill = PatternFill("solid", fgColor="1F4E78")
            header_fill = PatternFill("solid", fgColor="DCE6F1")
            for i, page in enumerate(pdf.pages, start=1):
                if progress_callback:
                    progress_callback(i, max(1, total_pages), f"Tablo aranıyor: Sayfa {i}/{total_pages}")

                ws = wb.create_sheet(_sanitize_sheet_title(f"Sayfa {i}"))
                ws.sheet_view.showGridLines = False
                ws.freeze_panes = "A2"

                tables = page.extract_tables(
                    table_settings={
                        "vertical_strategy": "lines",
                        "horizontal_strategy": "lines",
                        "snap_tolerance": 4,
                        "join_tolerance": 4,
                        "intersection_tolerance": 6,
                    }
                ) or []
                if not tables:
                    tables = page.extract_tables(
                        table_settings={
                            "vertical_strategy": "text",
                            "horizontal_strategy": "text",
                            "text_x_tolerance": 2,
                            "text_y_tolerance": 2,
                        }
                    ) or []
                current_row = 1

                cleaned_tables = []
                for table in tables:
                    cleaned_rows = []
                    for row in table or []:
                        cells = ["" if cell is None else str(cell).strip() for cell in row]
                        if any(cell for cell in cells):
                            cleaned_rows.append(cells)
                    if cleaned_rows:
                        cleaned_tables.append(cleaned_rows)

                if cleaned_tables:
                    for table_index, table_rows in enumerate(cleaned_tables, start=1):
                        title_cell = ws.cell(row=current_row, column=1, value=f"Tablo {table_index}")
                        title_cell.font = Font(bold=True, color="FFFFFF")
                        title_cell.fill = title_fill
                        title_cell.alignment = Alignment(horizontal="left", vertical="center")
                        current_row += 1
                        max_cols = max(len(r) for r in table_rows)
                        col_max = [0] * max_cols
                        for row_offset, row in enumerate(table_rows, start=0):
                            padded = row + [""] * (max_cols - len(row))
                            for col_index, value in enumerate(padded, start=1):
                                cell = ws.cell(row=current_row, column=col_index, value=value)
                                cell.border = border
                                cell.alignment = Alignment(vertical="center", horizontal="left", wrap_text=True)
                                if row_offset == 0:
                                    cell.font = Font(bold=True)
                                    cell.fill = header_fill
                                col_max[col_index - 1] = max(col_max[col_index - 1], len(str(value or "")))
                            ws.row_dimensions[current_row].height = 22
                            current_row += 1
                        for col_index, length in enumerate(col_max, start=1):
                            ws.column_dimensions[get_column_letter(col_index)].width = min(40, max(12, length + 2))
                        current_row += 2
                else:
                    text = (page.extract_text() or "").strip()
                    ws.cell(row=1, column=1, value="Bu sayfada tablo bulunamadı.")
                    if text:
                        ws.cell(row=3, column=1, value="Algılanan metin")
                        for row_index, line in enumerate([ln.strip() for ln in text.splitlines() if ln.strip()], start=4):
                            ws.cell(row=row_index, column=1, value=line)

        if not wb.sheetnames:
            ws = wb.create_sheet("Sayfa 1")
            ws.cell(row=1, column=1, value="İçerik bulunamadı.")

        wb.save(xlsx_path)
        return True
    except Exception as e:
        raise Exception(f"PDF tablo -> Excel Hatası: {e}") from e


def pdf_text_to_excel(
    pdf_path: str,
    xlsx_path: str,
    progress_callback=None,
    preserve_tables: bool = False,
    password: Optional[str] = None,
) -> bool:
    """
    PDF'i Excel'e aktarır.
    preserve_tables=True ise önce tablo yapısını korumaya çalışır.
    Aksi halde sayfa/satır bazlı metin aktarımı yapar.
    """
    if preserve_tables:
        try:
            return _pdf_tables_to_excel(pdf_path, xlsx_path, progress_callback=progress_callback, password=password)
        except Exception as e:
            raise Exception(f"PDF -> Excel (tablo koruma) Hatası: {e}") from e

    try:
        from openpyxl import Workbook
    except ImportError as e:
        raise Exception("PDF -> Excel için 'openpyxl' gerekli: python -m pip install openpyxl") from e

    try:
        reader = _open_pdf_reader(pdf_path, password=password, context="PDF -> Excel")
        num = len(reader.pages)
        wb = Workbook()
        ws = wb.active
        ws.title = "PDF Metni"
        ws.append(["Sayfa", "Satır No", "Metin"])

        for i in range(num):
            if progress_callback:
                progress_callback(i + 1, max(1, num), f"Sayfa {i + 1}/{num}")
            page = reader.pages[i]
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            if not lines:
                ws.append([i + 1, 1, "(Bu sayfada metin çıkarılamadı)"])
            else:
                for j, line in enumerate(lines, start=1):
                    ws.append([i + 1, j, line])

        wb.save(xlsx_path)
        return True
    except Exception as e:
        raise Exception(f"PDF -> Excel Hatası: {e}")


def _excel_to_pdf_win32com(xlsx_path: str, pdf_path: str) -> None:
    try:
        import pythoncom
        from win32com.client import DispatchEx
    except ImportError as e:
        raise ImportError("pywin32 gerekli.") from e

    pythoncom.CoInitialize()
    xl = None
    try:
        xl = DispatchEx("Excel.Application")
        xl.Visible = False
        try:
            xl.DisplayAlerts = False
        except Exception:
            pass
        wb = xl.Workbooks.Open(os.path.abspath(xlsx_path), ReadOnly=True)
        # 0 = xlTypePDF
        wb.ExportAsFixedFormat(0, os.path.abspath(pdf_path))
        wb.Close(SaveChanges=False)
    finally:
        if xl is not None:
            try:
                xl.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _excel_to_pdf_reportlab(xlsx_path: str, pdf_path: str) -> None:
    try:
        from openpyxl import load_workbook
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    except ImportError as e:
        raise ImportError("openpyxl ve reportlab gerekli.") from e

    wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    ws = wb.active
    rows: List[List[str]] = []
    max_cols = 0
    for row in ws.iter_rows(values_only=True):
        cells = ["" if c is None else str(c) for c in row]
        while cells and cells[-1] == "":
            cells.pop()
        if cells:
            max_cols = max(max_cols, len(cells))
            rows.append(cells)
    wb.close()

    if not rows:
        rows = [["(Boş sayfa)"]]
        max_cols = 1

    for r in rows:
        while len(r) < max_cols:
            r.append("")

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=landscape(A4),
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )
    usable_w = landscape(A4)[0] - 24 * mm
    col_w = usable_w / max(1, max_cols)
    t = Table(rows, colWidths=[col_w] * max_cols, repeatRows=1 if len(rows) > 1 else 0)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3a86ff")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
            ]
        )
    )
    doc.build([t])


# --- 6. EXCEL -> PDF ---
def excel_to_pdf(xlsx_path: str, pdf_path: str, progress_callback=None) -> bool:
    """
    Excel'i PDF yapar.
    Windows: Excel COM (Office). Diğer / yedek: openpyxl + reportlab (ilk sayfa, tablo görünümü).
    """
    if not os.path.isfile(xlsx_path):
        raise FileNotFoundError(f"Dosya bulunamadı: {xlsx_path}")
    ext = os.path.splitext(xlsx_path)[1].lower()
    if ext not in (".xlsx", ".xlsm", ".xltx", ".xltm"):
        raise ValueError("Desteklenen biçimler: .xlsx, .xlsm (eski .xls desteklenmez)")

    pdf_path = os.path.abspath(pdf_path)
    out_dir = os.path.dirname(pdf_path)
    if out_dir and not os.path.isdir(out_dir):
        raise FileNotFoundError(f"Hedef klasör bulunamadı: {out_dir}")

    if progress_callback:
        progress_callback(0, 2, "PDF oluşturuluyor...")

    if sys.platform == "win32":
        try:
            _excel_to_pdf_win32com(xlsx_path, pdf_path)
        except Exception:
            try:
                _excel_to_pdf_reportlab(xlsx_path, pdf_path)
            except Exception as e2:
                raise Exception(
                    "Excel -> PDF başarısız. Excel yüklüyse pywin32 kurun; değilse:\n"
                    "python -m pip install openpyxl reportlab\n\n"
                    f"Ayrıntı: {e2}"
                ) from e2
    else:
        try:
            _excel_to_pdf_reportlab(xlsx_path, pdf_path)
        except Exception as e:
            raise Exception(f"Excel -> PDF Hatası: {e}") from e

    if not os.path.isfile(pdf_path):
        raise Exception("PDF dosyası oluşmadı.")

    if progress_callback:
        progress_callback(2, 2, "Tamamlandı")
    return True


# --- 7. PDF SIKIŞTIRMA ---
def compress_pdf(input_path: str, output_path: str, progress_callback=None, password: Optional[str] = None) -> bool:
    """pikepdf ile akışları yeniden sıkıştırarak dosya boyutunu düşürmeye çalışır."""
    try:
        import pikepdf
    except ImportError as e:
        raise Exception("PDF sıkıştırma için 'pikepdf' gerekli.") from e

    try:
        if progress_callback:
            progress_callback(0, 2, "PDF yeniden paketleniyor...")
        if is_pdf_encrypted(input_path) and not password:
            raise Exception(f"PDF sıkıştırma için şifre gerekli: {os.path.basename(input_path)}")
        # pikepdf yeni sürümlerde password=None desteklenmez; boş dize kullanılmalı.
        open_password = (password or "").strip()
        target_path = output_path
        temp_output_path = None
        if os.path.abspath(input_path) == os.path.abspath(output_path):
            fd, temp_output_path = tempfile.mkstemp(suffix=".pdf", dir=os.path.dirname(output_path) or None)
            os.close(fd)
            target_path = temp_output_path
        with pikepdf.open(input_path, password=open_password) as pdf:
            pdf.save(
                target_path,
                compress_streams=True,
                recompress_flate=True,
                object_stream_mode=pikepdf.ObjectStreamMode.generate,
                linearize=True,
                stream_decode_level=pikepdf.StreamDecodeLevel.all,
            )
        if temp_output_path:
            os.replace(temp_output_path, output_path)
        # İkinci geçiş: PyMuPDF ile gömülü görselleri ve akışları yeniden sıkıştırır (genelde daha küçük dosya).
        try:
            import fitz

            out_dir = os.path.dirname(output_path) or None
            fd, tmp_mupdf = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
            os.close(fd)
            try:
                doc = fitz.open(output_path)
                try:
                    doc.save(
                        tmp_mupdf,
                        garbage=4,
                        clean=True,
                        deflate=True,
                        deflate_images=True,
                        deflate_fonts=True,
                        linear=True,
                        use_objstms=1,
                    )
                finally:
                    doc.close()
                old_sz = os.path.getsize(output_path)
                new_sz = os.path.getsize(tmp_mupdf)
                if new_sz < old_sz:
                    os.replace(tmp_mupdf, output_path)
                else:
                    try:
                        os.remove(tmp_mupdf)
                    except OSError:
                        pass
            except Exception:
                if os.path.isfile(tmp_mupdf):
                    try:
                        os.remove(tmp_mupdf)
                    except OSError:
                        pass
        except ImportError:
            pass
        except Exception:
            # pikepdf çıktısı yine de kullanılabilir
            pass
        if progress_callback:
            progress_callback(2, 2, "Tamamlandı")
        return True
    except Exception as e:
        raise Exception(f"PDF Sıkıştırma Hatası: {e}")
    finally:
        if "temp_output_path" in locals() and temp_output_path and os.path.exists(temp_output_path):
            try:
                os.remove(temp_output_path)
            except Exception:
                pass


# --- 8. PDF ŞİFRELEME ---
def encrypt_pdf(
    input_path: str,
    output_path: str,
    user_password: str,
    owner_password: Optional[str] = None,
    progress_callback=None,
    input_password: Optional[str] = None,
) -> bool:
    """Kullanıcı şifresi ile PDF şifreler (açmak için parola gerekir)."""
    try:
        import pikepdf
    except ImportError as e:
        raise Exception("PDF şifreleme için 'pikepdf' gerekli.") from e

    if not user_password:
        raise ValueError("Şifre boş olamaz.")

    try:
        if progress_callback:
            progress_callback(0, 2, "Şifre uygulanıyor...")
        owner = owner_password if owner_password else user_password
        if is_pdf_encrypted(input_path) and not input_password:
            raise Exception(f"PDF şifreleme için kaynak dosya şifresi gerekli: {os.path.basename(input_path)}")
        open_password = (input_password or "").strip()
        target_path = output_path
        temp_output_path = None
        if os.path.abspath(input_path) == os.path.abspath(output_path):
            fd, temp_output_path = tempfile.mkstemp(suffix=".pdf", dir=os.path.dirname(output_path) or None)
            os.close(fd)
            target_path = temp_output_path
        with pikepdf.open(input_path, password=open_password) as pdf:
            pdf.save(
                target_path,
                encryption=pikepdf.Encryption(
                    user=user_password,
                    owner=owner,
                    R=6,
                    allow=pikepdf.Permissions(extract=False),
                ),
            )
        if temp_output_path:
            os.replace(temp_output_path, output_path)
        if progress_callback:
            progress_callback(2, 2, "Tamamlandı")
        return True
    except Exception as e:
        raise Exception(f"PDF Şifreleme Hatası: {e}")
    finally:
        if "temp_output_path" in locals() and temp_output_path and os.path.exists(temp_output_path):
            try:
                os.remove(temp_output_path)
            except Exception:
                pass
